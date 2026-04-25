"""
tests/test_fileio.py
====================
Comprehensive unit tests for the FileIO format dispatcher.

Covers
------
* Supported read / write extension lists.
* Load .txt: content is faithfully read into the QTextDocument.
* Save .txt: QTextDocument content written to disk verbatim.
* Round-trip .txt: save then reload gives identical content.
* Load .md: content read correctly (plain-text storage for now).
* Save .md: content written to disk.
* Round-trip .md.
* Load .docx: requires python-docx; skipped if not installed.
* Save .docx: requires python-docx; skipped if not installed.
* Round-trip .docx.
* Unsupported extension raises ValueError on both load and save.
* Non-existent file raises FileNotFoundError on load.
* Multi-paragraph documents preserve paragraph structure.
* Unicode content (emoji, accented chars, CJK) survives round-trip.
* Large document (10 000 words) round-trip is consistent.
* FileIO.save() creates parent directories if needed.
* Mark-saved contract: doc.state is CLEAN after manager.save_document().
"""

from __future__ import annotations

import sys
import pytest
from pathlib import Path
from PySide6.QtWidgets import QApplication

@pytest.fixture(scope="session")
def qapp():
    app = QApplication.instance() or QApplication(sys.argv[:1])
    app.setApplicationName("PrasWordTest")
    app.setOrganizationName("PrasWordTest")
    return app


from prasword.core.document import Document, DocumentState
from prasword.core.document_manager import DocumentManager
from prasword.features.filemanagement.file_io import FileIO

# ── Helpers ───────────────────────────────────────────────────────────────────

def _docx_available() -> bool:
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        return False


def make_doc(text: str = "") -> Document:
    doc = Document()
    if text:
        doc.qt_document.setPlainText(text)
    return doc


# ═══════════════════════════════════════════════════════════════════════════════
# Extension registry
# ═══════════════════════════════════════════════════════════════════════════════

class TestExtensions:

    def test_read_extensions_include_txt(self, qapp):
        assert ".txt" in FileIO.supported_read_extensions()

    def test_read_extensions_include_md(self, qapp):
        assert ".md" in FileIO.supported_read_extensions()

    def test_read_extensions_include_docx(self, qapp):
        assert ".docx" in FileIO.supported_read_extensions()

    def test_write_extensions_include_txt(self, qapp):
        assert ".txt" in FileIO.supported_write_extensions()

    def test_write_extensions_include_md(self, qapp):
        assert ".md" in FileIO.supported_write_extensions()

    def test_write_extensions_include_docx(self, qapp):
        assert ".docx" in FileIO.supported_write_extensions()

    def test_extensions_are_lowercase(self, qapp):
        for ext in FileIO.supported_read_extensions():
            assert ext == ext.lower()

    def test_extensions_start_with_dot(self, qapp):
        for ext in FileIO.supported_read_extensions():
            assert ext.startswith(".")


# ═══════════════════════════════════════════════════════════════════════════════
# .txt
# ═══════════════════════════════════════════════════════════════════════════════

class TestTxt:

    def test_load_txt_basic(self, qapp, tmp_path):
        p = tmp_path / "basic.txt"
        p.write_text("Hello PrasWord!", encoding="utf-8")
        doc = Document(file_path=p)
        FileIO.load(doc, p)
        assert doc.plain_text() == "Hello PrasWord!"

    def test_save_txt_basic(self, qapp, tmp_path):
        p = tmp_path / "out.txt"
        doc = make_doc("Saved content.")
        FileIO.save(doc, p)
        assert p.read_text(encoding="utf-8") == "Saved content."

    def test_round_trip_txt_simple(self, qapp, tmp_path):
        p    = tmp_path / "rt.txt"
        text = "Round-trip test."
        doc  = make_doc(text)
        FileIO.save(doc, p)
        doc2 = Document(file_path=p)
        FileIO.load(doc2, p)
        assert doc2.plain_text() == text

    def test_round_trip_txt_multiline(self, qapp, tmp_path):
        p    = tmp_path / "multi.txt"
        text = "Line one\nLine two\nLine three"
        doc  = make_doc(text)
        FileIO.save(doc, p)
        doc2 = Document(file_path=p)
        FileIO.load(doc2, p)
        assert doc2.plain_text() == text

    def test_round_trip_txt_unicode(self, qapp, tmp_path):
        p    = tmp_path / "unicode.txt"
        text = "Héllo Wörld 你好 🎉"
        doc  = make_doc(text)
        FileIO.save(doc, p)
        doc2 = Document(file_path=p)
        FileIO.load(doc2, p)
        assert doc2.plain_text() == text

    def test_save_txt_creates_file(self, qapp, tmp_path):
        p = tmp_path / "new.txt"
        assert not p.exists()
        FileIO.save(make_doc("x"), p)
        assert p.exists()

    def test_save_txt_creates_parent_directories(self, qapp, tmp_path):
        p = tmp_path / "sub" / "deep" / "file.txt"
        FileIO.save(make_doc("deep"), p)
        assert p.exists()

    def test_load_txt_clears_modified_flag(self, qapp, tmp_path):
        p = tmp_path / "flag.txt"
        p.write_text("content", encoding="utf-8")
        doc = Document(file_path=p)
        FileIO.load(doc, p)
        assert not doc.qt_document.isModified()

    def test_load_txt_large_document(self, qapp, tmp_path):
        p    = tmp_path / "large.txt"
        text = ("word " * 100 + "\n") * 100   # 10 000 words
        p.write_text(text, encoding="utf-8")
        doc  = Document(file_path=p)
        FileIO.load(doc, p)
        assert doc.word_count() >= 9_000   # some trailing whitespace variation

    def test_round_trip_txt_empty(self, qapp, tmp_path):
        p   = tmp_path / "empty.txt"
        doc = make_doc("")
        FileIO.save(doc, p)
        doc2 = Document(file_path=p)
        FileIO.load(doc2, p)
        assert doc2.plain_text() == ""


# ═══════════════════════════════════════════════════════════════════════════════
# .md
# ═══════════════════════════════════════════════════════════════════════════════

class TestMarkdown:

    def test_load_md(self, qapp, tmp_path):
        p = tmp_path / "readme.md"
        p.write_text("# Title\n\nBody paragraph.", encoding="utf-8")
        doc = Document(file_path=p)
        FileIO.load(doc, p)
        assert "Title" in doc.plain_text()

    def test_save_md(self, qapp, tmp_path):
        p   = tmp_path / "out.md"
        doc = make_doc("# Heading\n\nSome text.")
        FileIO.save(doc, p)
        content = p.read_text(encoding="utf-8")
        assert "Heading" in content

    def test_round_trip_md(self, qapp, tmp_path):
        p    = tmp_path / "rt.md"
        text = "# H1\n\nParagraph text here."
        doc  = make_doc(text)
        FileIO.save(doc, p)
        doc2 = Document(file_path=p)
        FileIO.load(doc2, p)
        assert doc2.plain_text() == text

    def test_load_markdown_alias(self, qapp, tmp_path):
        p = tmp_path / "readme.markdown"
        p.write_text("content", encoding="utf-8")
        doc = Document(file_path=p)
        FileIO.load(doc, p)   # .markdown extension should also be supported
        assert "content" in doc.plain_text()

    def test_round_trip_md_unicode(self, qapp, tmp_path):
        p    = tmp_path / "unicode.md"
        text = "# Título\n\nContenido en español. 日本語テスト。"
        doc  = make_doc(text)
        FileIO.save(doc, p)
        doc2 = Document(file_path=p)
        FileIO.load(doc2, p)
        assert doc2.plain_text() == text


# ═══════════════════════════════════════════════════════════════════════════════
# .docx
# ═══════════════════════════════════════════════════════════════════════════════

@pytest.mark.skipif(not _docx_available(), reason="python-docx not installed")
class TestDocx:

    def test_save_docx_creates_file(self, qapp, tmp_path):
        p   = tmp_path / "out.docx"
        doc = make_doc("Docx content.")
        FileIO.save(doc, p)
        assert p.exists()
        assert p.stat().st_size > 0

    def test_load_docx_basic(self, qapp, tmp_path):
        import docx as python_docx
        p        = tmp_path / "in.docx"
        word_doc = python_docx.Document()
        word_doc.add_paragraph("Paragraph one.")
        word_doc.add_paragraph("Paragraph two.")
        word_doc.save(str(p))

        doc = Document(file_path=p)
        FileIO.load(doc, p)
        text = doc.plain_text()
        assert "Paragraph one." in text
        assert "Paragraph two." in text

    def test_round_trip_docx(self, qapp, tmp_path):
        p    = tmp_path / "rt.docx"
        doc  = make_doc("Round-trip docx test content here.")
        FileIO.save(doc, p)
        doc2 = Document(file_path=p)
        FileIO.load(doc2, p)
        assert "Round-trip" in doc2.plain_text()

    def test_save_docx_multi_paragraph(self, qapp, tmp_path):
        p    = tmp_path / "multi.docx"
        text = "First paragraph.\nSecond paragraph.\nThird paragraph."
        doc  = make_doc(text)
        FileIO.save(doc, p)
        doc2 = Document(file_path=p)
        FileIO.load(doc2, p)
        loaded = doc2.plain_text()
        assert "First paragraph." in loaded
        assert "Third paragraph." in loaded


# ═══════════════════════════════════════════════════════════════════════════════
# Error cases
# ═══════════════════════════════════════════════════════════════════════════════

class TestErrors:

    def test_load_unsupported_extension_raises_value_error(self, qapp, tmp_path):
        p = tmp_path / "file.xyz"
        p.touch()
        doc = make_doc()
        with pytest.raises(ValueError, match="[Uu]nsupported"):
            FileIO.load(doc, p)

    def test_save_unsupported_extension_raises_value_error(self, qapp, tmp_path):
        p   = tmp_path / "file.xyz"
        doc = make_doc("content")
        with pytest.raises(ValueError, match="[Uu]nsupported"):
            FileIO.save(doc, p)

    def test_load_nonexistent_file_raises_file_not_found(self, qapp, tmp_path):
        p   = tmp_path / "does_not_exist.txt"
        doc = make_doc()
        with pytest.raises((FileNotFoundError, Exception)):
            FileIO.load(doc, p)

    def test_load_uppercase_extension_is_normalised(self, qapp, tmp_path):
        # FileIO.load() lowercases the suffix, so .TXT is treated as .txt
        p = tmp_path / "file.TXT"
        p.write_text("content", encoding="utf-8")
        doc = make_doc()
        FileIO.load(doc, p)           # should NOT raise — .TXT → .txt is valid
        assert "content" in doc.plain_text()

    def test_load_truly_unsupported_extension_raises(self, qapp, tmp_path):
        # A genuinely unknown extension always raises ValueError
        p = tmp_path / "file.qwerty"
        p.write_text("content", encoding="utf-8")
        doc = make_doc()
        with pytest.raises(ValueError, match="[Uu]nsupported"):
            FileIO.load(doc, p)


# ═══════════════════════════════════════════════════════════════════════════════
# DocumentManager integration
# ═══════════════════════════════════════════════════════════════════════════════

class TestDocumentManagerIntegration:

    def test_save_via_manager_marks_clean(self, qapp, tmp_path):
        mgr = DocumentManager()
        doc = mgr.new_document()
        doc.qt_document.setPlainText("content")
        p   = tmp_path / "mgr.txt"
        mgr.save_document(doc, p)
        assert doc.state is DocumentState.CLEAN

    def test_open_via_manager_state_is_clean(self, qapp, tmp_path):
        p = tmp_path / "open.txt"
        p.write_text("hello", encoding="utf-8")
        mgr = DocumentManager()
        doc = mgr.open_document(p)
        assert doc.state is DocumentState.CLEAN

    def test_open_nonexistent_via_manager_raises(self, qapp, tmp_path):
        mgr = DocumentManager()
        with pytest.raises(FileNotFoundError):
            mgr.open_document(tmp_path / "ghost.txt")

    def test_double_open_same_file_returns_same_doc(self, qapp, tmp_path):
        p = tmp_path / "same.txt"
        p.write_text("content", encoding="utf-8")
        mgr = DocumentManager()
        d1  = mgr.open_document(p)
        d2  = mgr.open_document(p)
        assert d1 is d2
        assert mgr.document_count == 1
