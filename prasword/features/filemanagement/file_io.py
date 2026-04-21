"""
prasword.features.filemanagement.file_io
=========================================
File I/O adapters for every supported format.

``FileIO`` is a static dispatcher that routes load / save operations to the
correct format-specific handler based on the file extension.

Supported formats
-----------------
Read  : .docx, .md, .txt
Write : .docx, .md, .txt
Export: .pdf (via Qt's QPrinter)

Adding a new format
-------------------
1. Implement a class with ``load(doc, path)`` and/or ``save(doc, path)`` static methods.
2. Register it in ``_LOADERS`` / ``_WRITERS`` below.
"""

from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from PySide6.QtWidgets import QWidget
    from prasword.core.document import Document

from prasword.utils.logger import get_logger

log = get_logger(__name__)


# ---------------------------------------------------------------------------
# Format handlers
# ---------------------------------------------------------------------------

class _DocxHandler:
    """Read / write Microsoft Word (.docx) files via python-docx."""

    @staticmethod
    def load(doc: "Document", path: Path) -> None:
        try:
            import docx  # python-docx
        except ImportError:
            raise ImportError(
                "python-docx is required to open .docx files.\n"
                "  pip install python-docx"
            )

        word_doc = docx.Document(str(path))
        paragraphs = [p.text for p in word_doc.paragraphs]
        full_text = "\n".join(paragraphs)
        doc.qt_document.setPlainText(full_text)
        doc.qt_document.setModified(False)
        log.debug("Loaded .docx: %s (%d paragraphs)", path, len(paragraphs))

    @staticmethod
    def save(doc: "Document", path: Path) -> None:
        try:
            import docx
            from docx.shared import Pt
        except ImportError:
            raise ImportError(
                "python-docx is required to save .docx files.\n"
                "  pip install python-docx"
            )

        word_doc = docx.Document()
        # Write each Qt block as a paragraph.
        qt_doc = doc.qt_document
        block = qt_doc.begin()
        while block.isValid():
            word_doc.add_paragraph(block.text())
            block = block.next()

        word_doc.save(str(path))
        log.debug("Saved .docx: %s", path)


class _MarkdownHandler:
    """Read / write Markdown (.md) files."""

    @staticmethod
    def load(doc: "Document", path: Path) -> None:
        text = path.read_text(encoding="utf-8")
        doc.qt_document.setPlainText(text)
        doc.qt_document.setModified(False)
        log.debug("Loaded .md: %s", path)

    @staticmethod
    def save(doc: "Document", path: Path) -> None:
        # For now, save as plain text. A future enhancement can convert
        # QTextDocument formatting to proper Markdown syntax.
        text = doc.qt_document.toPlainText()
        path.write_text(text, encoding="utf-8")
        log.debug("Saved .md: %s", path)


class _PlainTextHandler:
    """Read / write plain-text (.txt) files."""

    @staticmethod
    def load(doc: "Document", path: Path) -> None:
        text = path.read_text(encoding="utf-8", errors="replace")
        doc.qt_document.setPlainText(text)
        doc.qt_document.setModified(False)
        log.debug("Loaded .txt: %s", path)

    @staticmethod
    def save(doc: "Document", path: Path) -> None:
        text = doc.qt_document.toPlainText()
        path.write_text(text, encoding="utf-8")
        log.debug("Saved .txt: %s", path)


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

_LOADERS = {
    ".docx": _DocxHandler,
    ".md":   _MarkdownHandler,
    ".markdown": _MarkdownHandler,
    ".txt":  _PlainTextHandler,
}

_WRITERS = {
    ".docx": _DocxHandler,
    ".md":   _MarkdownHandler,
    ".markdown": _MarkdownHandler,
    ".txt":  _PlainTextHandler,
}


class FileIO:
    """
    Static dispatcher for document I/O operations.

    Methods
    -------
    load(doc, path)
        Read a file from *path* into *doc*.
    save(doc, path)
        Write *doc* to *path*.
    export_pdf(doc, path, parent_widget)
        Export *doc* to PDF using Qt's printing subsystem.
    supported_read_extensions()
        List of supported input file extensions.
    supported_write_extensions()
        List of supported output file extensions.
    """

    @staticmethod
    def load(doc: "Document", path: Path) -> None:
        """
        Load a file into a Document.

        Parameters
        ----------
        doc : Document
        path : Path

        Raises
        ------
        ValueError
            If the file extension is not supported.
        """
        ext = path.suffix.lower()
        handler = _LOADERS.get(ext)
        if handler is None:
            raise ValueError(
                f"Unsupported file format: '{ext}'\n"
                f"Supported: {', '.join(_LOADERS)}"
            )
        handler.load(doc, path)

    @staticmethod
    def save(doc: "Document", path: Path) -> None:
        """
        Save a Document to disk.

        Parameters
        ----------
        doc : Document
        path : Path

        Raises
        ------
        ValueError
            If the file extension is not supported.
        """
        ext = path.suffix.lower()
        handler = _WRITERS.get(ext)
        if handler is None:
            raise ValueError(
                f"Unsupported file format for saving: '{ext}'\n"
                f"Supported: {', '.join(_WRITERS)}"
            )
        # Ensure the parent directory exists.
        path.parent.mkdir(parents=True, exist_ok=True)
        handler.save(doc, path)

    @staticmethod
    def export_pdf(
        doc: "Document",
        path: Path,
        parent_widget: "QWidget | None" = None,
    ) -> None:
        """
        Export a Document to PDF via Qt's printer subsystem.

        Parameters
        ----------
        doc : Document
        path : Path
            Output PDF file path.
        parent_widget : QWidget | None
            Optional parent for any error dialogs.

        Raises
        ------
        RuntimeError
            If PDF export fails.
        """
        from PySide6.QtPrintSupport import QPrinter
        from PySide6.QtGui import QTextDocument as _QTD

        printer = QPrinter(QPrinter.HighResolution)
        printer.setOutputFormat(QPrinter.PdfFormat)
        printer.setOutputFileName(str(path))
        printer.setPageSize(QPrinter.A4)

        # Print a *copy* of the QTextDocument to avoid modifying the original.
        qt_doc = _QTD()
        qt_doc.setHtml(doc.qt_document.toHtml())
        qt_doc.print_(printer)

        log.info("Exported PDF: %s", path)

    @staticmethod
    def supported_read_extensions() -> list[str]:
        """Return the list of loadable file extensions."""
        return list(_LOADERS.keys())

    @staticmethod
    def supported_write_extensions() -> list[str]:
        """Return the list of saveable file extensions."""
        return list(_WRITERS.keys())
